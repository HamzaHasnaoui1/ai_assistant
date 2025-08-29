#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to automatically generate Word reports from current_test_*.txt files
Generates a professional Word report with formatting for each test
"""

import os
import glob
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT

class WordReportGenerator:
    def __init__(self, output_dir="./results"):
        self.output_dir = output_dir
        self.ai_doc_dir = os.path.join(output_dir, "ai_assistant_documentation")
        
    def find_current_test_files(self):
        """Finds all current_test_*.txt files"""
        pattern = os.path.join(self.ai_doc_dir, "current_test_*.txt")
        return glob.glob(pattern)
    
    def parse_test_file(self, file_path):
        """Parses the content of a test file and extracts information"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract test file name
            test_name = os.path.basename(file_path)
            
            # Extract interactions using a different approach
            interactions = []
            
            # Look for all occurrences of "AI ASSISTANT INTERACTION DOCUMENTATION"
            import re
            pattern = r'AI ASSISTANT INTERACTION DOCUMENTATION(.*?)(?=AI ASSISTANT INTERACTION DOCUMENTATION|$)'
            matches = re.findall(pattern, content, re.DOTALL)
            
            for match in matches:
                # Add section title at the beginning
                section = "AI ASSISTANT INTERACTION DOCUMENTATION" + match
                # Extract interaction information
                interaction = self.extract_interaction_info(section)
                if interaction:
                    interactions.append(interaction)
            
            # If no interactions found with the new method, try the old one
            if not interactions:
                print("⚠️ New parsing method failed, trying with old method...")
                sections = content.split("========================================")
                
                for section in sections:
                    if "AI ASSISTANT INTERACTION DOCUMENTATION" in section and len(section.strip()) > 100:
                        # Extract interaction information
                        interaction = self.extract_interaction_info(section)
                        if interaction:
                            interactions.append(interaction)
            
            return {
                'test_name': test_name,
                'file_path': file_path,
                'interactions': interactions,
                'total_interactions': len(interactions)
            }
            
        except Exception as e:
            print(f"Error parsing {file_path}: {e}")
            return None
    
    def extract_interaction_info(self, section):
        """Extracts interaction information from a section"""
        try:
            # Extract timestamp
            timestamp_match = re.search(r'Timestamp: (.+)', section)
            timestamp = timestamp_match.group(1).strip() if timestamp_match else "N/A"
            
            # Extract question
            question_match = re.search(r'Question Asked:\s*\n"([^"]+)"', section, re.DOTALL)
            question = question_match.group(1).strip() if question_match else "N/A"
            
            # Extract answer
            answer_match = re.search(r'AI Response:\s*\n"([^"]+)"', section, re.DOTALL)
            answer = answer_match.group(1).strip() if answer_match else "N/A"
            
            # Extract category
            category_match = re.search(r'Response Category: (.+)', section)
            category = category_match.group(1).strip() if category_match else "N/A"
            
            # Extract length
            length_match = re.search(r'Response Length: (\d+) characters', section)
            length = length_match.group(1) if length_match else "N/A"
            
            # Extract word count
            words_match = re.search(r'Word Count: (\d+)', section)
            words = words_match.group(1) if words_match else "N/A"
            
            # Extract quality score
            quality_match = re.search(r'Quality Score: ([\d.]+)%', section)
            quality = quality_match.group(1) if quality_match else "N/A"
            
            # Extract question mark presence
            question_mark_match = re.search(r'Question Mark at End: (.+)', section)
            question_mark = question_mark_match.group(1).strip() if question_mark_match else "N/A"
            
            # Extract screenshot paths
            q_screenshot_match = re.search(r'Question Screenshot: (.+)', section)
            q_screenshot = q_screenshot_match.group(1).strip() if q_screenshot_match else "N/A"
            
            r_screenshot_match = re.search(r'Response Screenshot: (.+)', section)
            r_screenshot = r_screenshot_match.group(1).strip() if r_screenshot_match else "N/A"
            
            return {
                'timestamp': timestamp,
                'question': question,
                'answer': answer,
                'category': category,
                'length': length,
                'words': words,
                'quality': quality,
                'question_mark': question_mark,
                'q_screenshot': q_screenshot,
                'r_screenshot': r_screenshot
            }
            
        except Exception as e:
            print(f"Error extracting information: {e}")
            return None
    
    def create_word_document(self, test_data):
        """Creates a Word document for a given test"""
        try:
            # Create a new document
            doc = Document()
            
            # Main title
            title = doc.add_heading('Test Report - AI Assistant', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add company logo or image (optional)
            # doc.add_picture('logo_atg.png', width=Inches(2))
            
            # Test information
            doc.add_heading('General Information', level=1)
            
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            
            # Headers
            header_cells = info_table.rows[0].cells
            header_cells[0].text = 'Property'
            header_cells[1].text = 'Value'
            
            # Information
            info_data = [
                ('Test Name', test_data['test_name']),
                ('Source File', test_data['file_path']),
                ('Total Interactions', str(test_data['total_interactions'])),
                ('Generation Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            ]
            
            for prop, value in info_data:
                row_cells = info_table.add_row().cells
                row_cells[0].text = prop
                row_cells[1].text = value
            
            # Calculation Methods
            doc.add_heading('Calculation Methods', level=1)
            
            # Category calculation explanation
            doc.add_heading('Response Category Classification', level=2)
            category_explanation = doc.add_paragraph()
            category_explanation.add_run('The AI responses are automatically classified into three categories based on length and word count thresholds:').bold = True
            doc.add_paragraph()
            
            category_table = doc.add_table(rows=1, cols=3)
            category_table.style = 'Table Grid'
            category_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Category headers
            cat_headers = category_table.rows[0].cells
            cat_headers[0].text = 'Category'
            cat_headers[1].text = 'Character Threshold'
            cat_headers[2].text = 'Word Threshold'
            
            # Category data
            category_data = [
                ('LONG', '> 250 characters', '> 50 words'),
                ('MEDIUM', '≥ 100 characters', '≥ 20 words'),
                ('SHORT', '< 100 characters', '< 20 words')
            ]
            
            for cat, char_thresh, word_thresh in category_data:
                row_cells = category_table.add_row().cells
                row_cells[0].text = cat
                row_cells[1].text = char_thresh
                row_cells[2].text = word_thresh
            
            doc.add_paragraph()
            doc.add_paragraph('Note: A response is classified as LONG if it exceeds EITHER the character OR word threshold. MEDIUM requires meeting BOTH thresholds. SHORT applies when BOTH thresholds are below the minimum.')
            
            # Quality score calculation explanation
            doc.add_heading('Quality Score Calculation', level=2)
            quality_explanation = doc.add_paragraph()
            quality_explanation.add_run('The quality score is calculated based on the response length relative to the LONG threshold:').bold = True
            doc.add_paragraph()
            
            quality_formula = doc.add_paragraph()
            quality_formula.add_run('Quality Score = min(100, (Response Length / 250) × 100)').italic = True
            doc.add_paragraph()
            
            quality_details = doc.add_paragraph()
            quality_details.add_run('Where:').bold = True
            doc.add_paragraph('• Response Length = Number of characters in the AI response')
            doc.add_paragraph('• 250 = LONG_CHAR_THRESHOLD (maximum threshold for LONG category)')
            doc.add_paragraph('• The score is capped at 100% maximum')
            doc.add_paragraph('• Higher scores indicate responses that are closer to or exceed the LONG threshold')
            
            doc.add_paragraph()
            
            # Interactions summary
            doc.add_heading('Interactions Summary', level=1)
            
            if test_data['interactions']:
                summary_table = doc.add_table(rows=1, cols=6)
                summary_table.style = 'Table Grid'
                
                # Summary headers
                summary_headers = summary_table.rows[0].cells
                summary_headers[0].text = 'No.'
                summary_headers[1].text = 'Timestamp'
                summary_headers[2].text = 'Question'
                summary_headers[3].text = 'Category'
                summary_headers[4].text = 'Quality Score'
                summary_headers[5].text = 'Question Mark'
                
                # Summary data
                for i, interaction in enumerate(test_data['interactions'], 1):
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = interaction['timestamp']
                    row_cells[2].text = interaction['question'][:50] + "..." if len(interaction['question']) > 50 else interaction['question']
                    row_cells[3].text = interaction['category']
                    row_cells[4].text = f"{interaction['quality']}%"
                    row_cells[5].text = interaction['question_mark']
            
            # Interaction details
            doc.add_heading('Interaction Details', level=1)
            
            for i, interaction in enumerate(test_data['interactions'], 1):
                # Interaction title
                doc.add_heading(f'Interaction {i}', level=2)
                
                # Question section
                doc.add_heading('Question', level=3)
                doc.add_paragraph(interaction['question'])
                
                # AI Response section
                doc.add_heading('AI Response', level=3)
                doc.add_paragraph(interaction['answer'])
                
                # Metrics table
                metrics_table = doc.add_table(rows=1, cols=2)
                metrics_table.style = 'Table Grid'
                metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Headers
                metrics_headers = metrics_table.rows[0].cells
                metrics_headers[0].text = 'Metric'
                metrics_headers[1].text = 'Value'
                
                # Metrics
                metrics_data = [
                    ('Timestamp', interaction['timestamp']),
                    ('Category', interaction['category']),
                    ('Length', f"{interaction['length']} characters"),
                    ('Words', interaction['words']),
                    ('Quality Score', f"{interaction['quality']}%"),
                    ('Question Mark', interaction['question_mark'])
                ]
                
                for prop, value in metrics_data:
                    row_cells = metrics_table.add_row().cells
                    row_cells[0].text = prop
                    row_cells[1].text = value
                
                # Screenshots section
                doc.add_heading('Screenshots', level=3)
                
                # Screenshots with error handling
                self.add_screenshots_to_document(doc, interaction, i)
                
                # Add space between interactions
                doc.add_paragraph()
            
            # Detailed statistics
            if test_data['interactions']:
                doc.add_heading('Detailed Statistics', level=1)
                
                # Calculate statistics
                categories = [i['category'] for i in test_data['interactions']]
                qualities = [float(i['quality']) for i in test_data['interactions'] if i['quality'] != 'N/A']
                lengths = [int(i['length']) for i in test_data['interactions'] if i['length'] != 'N/A']
                words = [int(i['words']) for i in test_data['interactions'] if i['words'] != 'N/A']
                question_marks = [i['question_mark'] for i in test_data['interactions']]
                
                # Main statistics
                main_stats_data = [
                    ('Total Interactions', str(len(test_data['interactions']))),
                    ('Present Categories', ', '.join(set(categories))),
                    ('Average Quality Score', f"{sum(qualities)/len(qualities):.2f}%" if qualities else "N/A"),
                    ('Average Length', f"{sum(lengths)/len(lengths):.0f} characters" if lengths else "N/A"),
                    ('Average Words', f"{sum(words)/len(words):.0f} words" if words else "N/A"),
                    ('Responses with Question Mark', f"{question_marks.count('YES')}/{len(question_marks)} ({question_marks.count('YES')/len(question_marks)*100:.1f}%)" if question_marks else "N/A")
                ]
                
                main_stats_table = doc.add_table(rows=1, cols=2)
                main_stats_table.style = 'Table Grid'
                main_stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Headers
                main_stats_headers = main_stats_table.rows[0].cells
                main_stats_headers[0].text = 'Metric'
                main_stats_headers[1].text = 'Value'
                
                # Main statistics
                for metric, value in main_stats_data:
                    row_cells = main_stats_table.add_row().cells
                    row_cells[0].text = metric
                    row_cells[1].text = value
                
                doc.add_paragraph()
                
                # Extreme statistics
                if qualities:
                    extreme_stats_data = [
                        ('Maximum Quality Score', f"{max(qualities):.2f}%"),
                        ('Minimum Quality Score', f"{min(qualities):.2f}%"),
                        ('Quality Range', f"{max(qualities) - min(qualities):.2f}%")
                    ]
                    
                    extreme_stats_table = doc.add_table(rows=1, cols=2)
                    extreme_stats_table.style = 'Table Grid'
                    extreme_stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers
                    extreme_headers = extreme_stats_table.rows[0].cells
                    extreme_headers[0].text = 'Metric'
                    extreme_headers[1].text = 'Value'
                    
                    # Extreme statistics
                    for metric, value in extreme_stats_data:
                        row_cells = extreme_stats_table.add_row().cells
                        row_cells[0].text = metric
                        row_cells[1].text = value
                
                doc.add_paragraph()
                
                # Category analysis
                category_analysis = {}
                for interaction in test_data['interactions']:
                    cat = interaction['category']
                    if cat not in category_analysis:
                        category_analysis[cat] = {'count': 0, 'qualities': [], 'lengths': []}
                    
                    category_analysis[cat]['count'] += 1
                    if interaction['quality'] != 'N/A':
                        category_analysis[cat]['qualities'].append(float(interaction['quality']))
                    if interaction['length'] != 'N/A':
                        category_analysis[cat]['lengths'].append(int(interaction['length']))
                
                if category_analysis:
                    doc.add_heading('Category Analysis', level=2)
                    
                    cat_table = doc.add_table(rows=1, cols=4)
                    cat_table.style = 'Table Grid'
                    cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers
                    cat_headers = cat_table.rows[0].cells
                    cat_headers[0].text = 'Category'
                    cat_headers[1].text = 'Count'
                    cat_headers[2].text = 'Average Quality'
                    cat_headers[3].text = 'Average Length'
                    
                    # Data by category
                    for cat, data in category_analysis.items():
                        row_cells = cat_table.add_row().cells
                        row_cells[0].text = cat
                        row_cells[1].text = str(data['count'])
                        
                        if data['qualities']:
                            avg_quality = sum(data['qualities']) / len(data['qualities'])
                            row_cells[2].text = f"{avg_quality:.2f}%"
                        else:
                            row_cells[2].text = "N/A"
                        
                        if data['lengths']:
                            avg_length = sum(data['lengths']) / len(data['lengths'])
                            row_cells[3].text = f"{avg_length:.0f} chars"
                        else:
                            row_cells[3].text = "N/A"
            
            # Screenshots gallery section removed as requested
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph('Report automatically generated by ATG test system')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return doc
            
        except Exception as e:
            print(f"Error creating Word document: {e}")
            return None
    
    def add_screenshots_to_document(self, doc, interaction, interaction_number):
        """Adds screenshots to the Word document"""
        try:
            # Question screenshots
            if interaction['q_screenshot'] != 'N/A':
                q_screenshot_path = interaction['q_screenshot']
                if os.path.exists(q_screenshot_path):
                    doc.add_paragraph(f'Question {interaction_number}:')
                    try:
                        # Insert image with appropriate size
                        doc.add_picture(q_screenshot_path, width=Inches(6))
                        doc.add_paragraph(f'File: {os.path.basename(q_screenshot_path)}')
                    except Exception as e:
                        doc.add_paragraph(f'[ERROR] Error inserting image: {e}')
                        doc.add_paragraph(f'Path: {q_screenshot_path}')
                else:
                    doc.add_paragraph(f'[INFO] Question screenshot not available: {os.path.basename(q_screenshot_path)}')
                    doc.add_paragraph('Note: Screenshot file was not found during test execution')
            else:
                doc.add_paragraph('No question screenshot available')
            
            # Space between screenshots
            doc.add_paragraph()
            
            # Response screenshots
            if interaction['r_screenshot'] != 'N/A':
                r_screenshot_path = interaction['r_screenshot']
                if os.path.exists(r_screenshot_path):
                    doc.add_paragraph(f'Response {interaction_number}:')
                    try:
                        # Insert image with appropriate size
                        doc.add_picture(r_screenshot_path, width=Inches(6))
                        doc.add_paragraph(f'File: {os.path.basename(r_screenshot_path)}')
                    except Exception as e:
                        doc.add_paragraph(f'[ERROR] Error inserting image: {e}')
                        doc.add_paragraph(f'Path: {r_screenshot_path}')
                else:
                    doc.add_paragraph(f'[INFO] Response screenshot not available: {os.path.basename(r_screenshot_path)}')
                    doc.add_paragraph('Note: Screenshot file was not found during test execution')
            else:
                doc.add_paragraph('No response screenshot available')
            
            # Space after screenshots
            doc.add_paragraph()
            
        except Exception as e:
            doc.add_paragraph(f'[ERROR] Error adding screenshots: {e}')
    
    # add_summary_screenshots function removed as requested
    
    def save_word_document(self, doc, test_data):
        """Saves the Word document"""
        try:
            # Create Word filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            word_filename = f"test_report_{timestamp}.docx"
            word_path = os.path.join(self.ai_doc_dir, word_filename)
            
            # Save document
            doc.save(word_path)
            print(f"Word report saved: {word_path}")
            
            return word_path
            
        except Exception as e:
            print(f"Error saving Word document: {e}")
            return None
    
    def generate_all_reports(self):
        """Generates Word reports for all found tests"""
        print("Searching for test files...")
        test_files = self.find_current_test_files()
        
        if not test_files:
            print("No test files found.")
            return
        
        print(f"Found {len(test_files)} test file(s).")
        
        for test_file in test_files:
            print(f"\nProcessing: {test_file}")
            
            # Parse test file
            test_data = self.parse_test_file(test_file)
            if not test_data:
                print(f"Cannot parse {test_file}")
                continue
            
            # Create Word document
            doc = self.create_word_document(test_data)
            if not doc:
                print(f"Cannot create Word document for {test_file}")
                continue
            
            # Save document
            word_path = self.save_word_document(doc, test_data)
            if word_path:
                print(f"Word report generated successfully: {word_path}")
            else:
                print(f"Failed to save Word report for {test_file}")
    
    def generate_report_for_latest_test(self):
        """Generates a Word report for the most recent test"""
        test_files = self.find_current_test_files()
        
        if not test_files:
            print("No test files found.")
            return None
        
        # Sort by modification date (most recent first)
        test_files.sort(key=lambda x: os.path.getmtime(x), reverse=True)
        latest_test = test_files[0]
        
        print(f"Generating report for most recent test: {latest_test}")
        
        # Parse test file
        test_data = self.parse_test_file(latest_test)
        if not test_data:
            print(f"Cannot parse {latest_test}")
            return None
        
        # Create Word document
        doc = self.create_word_document(test_data)
        if not doc:
            print(f"Cannot create Word document for {latest_test}")
            return None
        
        # Save document
        word_path = self.save_word_document(doc, test_data)
        if word_path:
            print(f"Word report generated successfully: {word_path}")
            return word_path
        else:
            print(f"Failed to save Word report for {latest_test}")
            return None

def main():
    """Main function"""
    print("=== ATG Test Word Report Generator ===")
    
    # Create generator
    generator = WordReportGenerator()
    
    # Check if directory exists
    if not os.path.exists(generator.ai_doc_dir):
        print(f"Directory {generator.ai_doc_dir} does not exist.")
        return
    
    # Generate report for most recent test
    print("\nGenerating report for most recent test...")
    word_path = generator.generate_report_for_latest_test()
    
    if word_path:
        print(f"\n[SUCCESS] Word report generated successfully: {word_path}")
    else:
        print("\n[ERROR] Failed to generate Word report")

if __name__ == "__main__":
    main()
